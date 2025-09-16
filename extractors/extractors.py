
import pdfplumber
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import docx
import pandas as pd
import ezdxf
from typing import Dict, Any
import json
import os

class BaseExtractor:
    """Base class for all extractors"""

    def extract(self, filepath: str) -> Dict[str, Any]:
        """Extract text and metadata from file"""
        raise NotImplementedError

class PDFExtractor(BaseExtractor):
    """Extract text from PDF files using pdfplumber with OCR fallback"""

    def extract(self, filepath: str) -> Dict[str, Any]:
        result = {
            'text': '',
            'metadata': {
                'pages': 0,
                'extraction_method': [],
                'tables_found': 0
            }
        }

        try:
            with pdfplumber.open(filepath) as pdf:
                result['metadata']['pages'] = len(pdf.pages)
                all_text = []
                tables_count = 0

                for i, page in enumerate(pdf.pages):
                    # Try direct text extraction first
                    text = page.extract_text()

                    if text and len(text.strip()) > 30:
                        all_text.append(text)
                        result['metadata']['extraction_method'].append(f'page_{i+1}_direct')
                    else:
                        # Fallback to OCR
                        try:
                            img_pages = convert_from_path(filepath, first_page=i+1, last_page=i+1)
                            for img in img_pages:
                                ocr_text = pytesseract.image_to_string(img, lang='eng+mal')
                                all_text.append(ocr_text)
                                result['metadata']['extraction_method'].append(f'page_{i+1}_ocr')
                        except Exception as e:
                            print(f"OCR failed for page {i+1}: {e}")

                    # Extract tables
                    try:
                        tables = page.extract_tables()
                        if tables:
                            tables_count += len(tables)
                            for table in tables:
                                table_text = '\n'.join([' | '.join([str(cell) for cell in row if cell]) for row in table])
                                all_text.append(f"TABLE:\n{table_text}")
                    except:
                        pass

                result['text'] = '\n\n'.join(all_text)
                result['metadata']['tables_found'] = tables_count

        except Exception as e:
            result['metadata']['error'] = str(e)

        return result

class ImageExtractor(BaseExtractor):
    """Extract text from images using OCR"""

    def extract(self, filepath: str) -> Dict[str, Any]:
        result = {
            'text': '',
            'metadata': {
                'ocr_method': 'tesseract',
                'image_size': None
            }
        }

        try:
            img = Image.open(filepath)
            result['metadata']['image_size'] = img.size

            # Try OCR with multiple languages
            text = pytesseract.image_to_string(img, lang='eng+mal')
            result['text'] = text

        except Exception as e:
            result['metadata']['error'] = str(e)

        return result

class DOCXExtractor(BaseExtractor):
    """Extract text from DOCX files"""

    def extract(self, filepath: str) -> Dict[str, Any]:
        result = {
            'text': '',
            'metadata': {
                'paragraphs': 0,
                'tables': 0
            }
        }

        try:
            doc = docx.Document(filepath)
            all_text = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    all_text.append(paragraph.text)

            result['metadata']['paragraphs'] = len([p for p in doc.paragraphs if p.text.strip()])

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(' | '.join(row_text))
                all_text.append('TABLE:\n' + '\n'.join(table_text))

            result['metadata']['tables'] = len(doc.tables)
            result['text'] = '\n\n'.join(all_text)

        except Exception as e:
            result['metadata']['error'] = str(e)

        return result

class CADExtractor(BaseExtractor):
    """Extract text from CAD files (DXF)"""

    def extract(self, filepath: str) -> Dict[str, Any]:
        result = {
            'text': '',
            'metadata': {
                'entities': 0,
                'text_entities': 0,
                'layers': []
            }
        }

        try:
            if filepath.lower().endswith('.dxf'):
                doc = ezdxf.readfile(filepath)
                all_text = []
                text_entities = 0

                # Get all layers
                result['metadata']['layers'] = [layer.dxf.name for layer in doc.layers]

                # Extract text entities
                msp = doc.modelspace()
                for entity in msp:
                    result['metadata']['entities'] += 1

                    if entity.dxftype() in ['TEXT', 'MTEXT']:
                        all_text.append(entity.dxf.text if hasattr(entity.dxf, 'text') else str(entity))
                        text_entities += 1
                    elif entity.dxftype() == 'DIMENSION':
                        # Extract dimension text
                        if hasattr(entity.dxf, 'text'):
                            all_text.append(entity.dxf.text)
                            text_entities += 1

                result['metadata']['text_entities'] = text_entities
                result['text'] = '\n'.join(all_text)

        except Exception as e:
            result['metadata']['error'] = str(e)
            # For DWG files or if ezdxf fails, return metadata only
            result['text'] = f"CAD file detected: {os.path.basename(filepath)}"

        return result

class CSVExtractor(BaseExtractor):
    """Extract content from CSV files"""

    def extract(self, filepath: str) -> Dict[str, Any]:
        result = {
            'text': '',
            'metadata': {
                'rows': 0,
                'columns': 0,
                'encoding': 'utf-8'
            }
        }

        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    df = pd.read_csv(filepath, encoding=encoding)
                    result['metadata']['encoding'] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode CSV with any common encoding")

            result['metadata']['rows'] = len(df)
            result['metadata']['columns'] = len(df.columns)

            # Convert to text representation
            text_parts = [
                f"CSV Table with {len(df)} rows and {len(df.columns)} columns:",
                f"Columns: {', '.join(df.columns)}",
                "Data preview:",
                df.head(10).to_string(index=False)
            ]

            result['text'] = '\n\n'.join(text_parts)

        except Exception as e:
            result['metadata']['error'] = str(e)

        return result

class TXTExtractor(BaseExtractor):
    """Extract content from plain text files"""

    def extract(self, filepath: str) -> Dict[str, Any]:
        result = {
            'text': '',
            'metadata': {
                'encoding': 'utf-8',
                'lines': 0
            }
        }

        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(filepath, 'r', encoding=encoding) as f:
                        text = f.read()
                    result['metadata']['encoding'] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode text file with any common encoding")

            result['text'] = text
            result['metadata']['lines'] = len(text.split('\n'))

        except Exception as e:
            result['metadata']['error'] = str(e)

        return result

class ExtractorFactory:
    """Factory to get the appropriate extractor for each file type"""

    extractors = {
        'PDF': PDFExtractor,
        'IMAGE': ImageExtractor,
        'DOCX': DOCXExtractor,
        'DOC': DOCXExtractor,  # Use same as DOCX for now
        'DXF': CADExtractor,
        'DWG': CADExtractor,
        'CSV': CSVExtractor,
        'TXT': TXTExtractor,
        'XLSX': CSVExtractor,  # Basic extraction, could be enhanced
    }

    @classmethod
    def get_extractor(cls, file_type: str) -> BaseExtractor:
        """Get appropriate extractor for file type"""
        extractor_class = cls.extractors.get(file_type)
        if extractor_class:
            return extractor_class()
        else:
            raise ValueError(f"No extractor available for file type: {file_type}")
